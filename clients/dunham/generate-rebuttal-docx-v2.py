#!/usr/bin/env python3
"""Generate Expert Rebuttal Report V2 as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading 1 style (section headers)
h1 = doc.styles['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Pt(14)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(8)
h1.paragraph_format.keep_with_next = True

# Heading 2 style (subsection headers)
h2 = doc.styles['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Pt(12)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
h2.paragraph_format.space_before = Pt(16)
h2.paragraph_format.space_after = Pt(6)
h2.paragraph_format.keep_with_next = True


def add_bold_text(paragraph, bold_text, normal_text=""):
    run = paragraph.add_run(bold_text)
    run.bold = True
    if normal_text:
        paragraph.add_run(normal_text)


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_numbered_item(doc, bold_text, normal_text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(bold_text)
    run.bold = True
    p.add_run(normal_text)
    return p


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_key_finding(doc, text):
    """Add a bordered key finding box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    # Add border via XML
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), '1a1a1a')
        pBdr.append(border)
    pPr.append(pBdr)
    # Add indentation for padding effect
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '144')
    ind.set(qn('w:right'), '144')
    pPr.append(ind)
    run = p.add_run("Finding: ")
    run.bold = True
    p.add_run(text)
    return p


def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    # Set font size in header
    for paragraph in table.rows[0].cells[0].paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
    return table


# ============================================================
# HEADER / TITLE PAGE
# ============================================================

# Logo
logo_path = os.path.join(os.path.dirname(__file__), '..', 'assets', 'imgs', 'logos', 'hyder-media-logo.png')
if os.path.exists(logo_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(0.6))

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Expert Rebuttal Report')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Response to Expert Report of Caryn Brown, Digital Media Butterfly\ndated February 6, 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)

# Separator
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('_' * 60)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

# Meta info
meta_items = [
    ("Prepared for: ", "Dunham Law Firm, P.C. / Dunham & Jones L.L.P."),
    ("Prepared by: ", "Kenny Hyder, Founder, Hyder Media"),
    ("Case: ", "Dunnam & Dunnam LLP v. Dunham Law Firm, P.C., et al."),
    ("", "Civil Action No. 6:21-cv-1041-ADA-DGT, W.D. Texas, Waco Division"),
    ("Date: ", "March 1, 2026"),
]
for bold_part, normal_part in meta_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    if bold_part:
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(10)
    run = p.add_run(normal_part)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)

doc.add_page_break()

# ============================================================
# I. INTRODUCTION
# ============================================================

doc.add_heading('I. Introduction and Qualifications', level=1)

add_para(doc, 'I have been retained by counsel for Defendants to review and respond to the Expert Report of Caryn Brown (\u201cBrown Report\u201d), dated February 6, 2026, prepared on behalf of Plaintiff Dunnam & Dunnam LLP. My analysis addresses the technical accuracy, methodological rigor, and factual basis of Brown\u2019s claims regarding Defendants\u2019 online properties and their relationship to the Agreed Judgment dated April 24, 1998 (the \u201c1998 Judgment\u201d).')

add_para(doc, 'I am a digital marketing and search advertising consultant with over 20 years of hands-on experience in search engine optimization (SEO), paid search advertising (Google Ads/PPC), content marketing, online reputation management, and competitive analysis. I have directed search marketing strategy for multiple Fortune 500 companies, including Fortune 20 and Fortune 5 businesses. I am a recognized contributing author at Search Engine Journal, a veteran speaker at major industry conferences including Pubcon, SMX, Affiliate Summit, and the Direct Marketing Association, and I serve as a judge for the US Search Awards. My full curriculum vitae is attached.')

add_para(doc, 'My opinions are based on my professional experience, direct observation using industry-standard tools and methodologies, and review of the Brown Report, the 1998 Agreed Judgment, and the publicly observable online properties at issue. I am being compensated at a rate of $300 per hour for my work in this matter. My compensation is not contingent on the outcome of this case or the opinions expressed herein.')

add_para(doc, 'Since my initial review of the Brown Report, I have conducted independent analysis using the following industry-standard tools, all of which generate verifiable, downloadable reports:')

add_bullet(doc, ' \u2014 full site crawl of DunhamLaw.com (2,104 internal URLs, 1,548 HTML pages) with custom extraction configured to search for the keyword \u201cdunnam\u201d in HTML source, rendered page text, and content areas', 'Screaming Frog SEO Spider')
add_bullet(doc, ' \u2014 paid and organic keyword analysis for DunhamLaw.com, DunhamJones.com, and DunnamLaw.com (February 2025 through January 2026)', 'SimilarWeb')
add_bullet(doc, ' \u2014 organic position data for all three domains, plus backlink and anchor text analysis for DunhamLaw.com', 'SEMrush')
add_bullet(doc, ' \u2014 review of all paid advertisements associated with the DunhamLaw.com domain', 'Google Ads Transparency Center')

add_para(doc, 'The raw exports from each of these tools are available upon request and form the evidentiary basis for the findings presented in this report.')

# ============================================================
# II. SUMMARY
# ============================================================

doc.add_heading('II. Summary of Opinions', level=1)

add_para(doc, 'After thorough review of the Brown Report, the underlying court orders, and independent analysis of Defendants\u2019 online properties using industry-standard tools, it is my opinion that:')

add_numbered_item(doc,
    'The Brown Report fails to identify a single concrete violation of the 1998 Judgment, and my independent analysis confirms no such violation exists. ',
    'A full-site crawl of DunhamLaw.com\u20142,104 URLs, 1,548 HTML pages\u2014found zero instances of the word \u201cdunnam\u201d anywhere on the site. The Google Ads Transparency Center confirms all Defendants\u2019 advertisements run under the name \u201cDunham & Jones Attorneys at Law, P.C.\u201d\u2014not any prohibited name. Third-party keyword data from SimilarWeb and SEMrush confirms Defendants have zero paid keywords targeting \u201cDunnam\u201d and receive zero click traffic from organic \u201cdunnam\u201d queries.')

add_numbered_item(doc,
    'The Brown Report fundamentally mischaracterizes the scope of the 1998 Judgment. ',
    'The Judgment restricts the use of specific name formulations (e.g., \u201cDunham Law Firm,\u201d \u201cDunham, Attorney at Law\u201d) in specific Texas counties. It does not prohibit Defendants from using their legal surname \u201cDunham,\u201d from operating the domain DunhamLaw.com, from practicing law in those counties, or from advertising under non-prohibited names such as \u201cDunham & Jones.\u201d')

add_numbered_item(doc,
    'Brown\u2019s methodology is insufficient, undocumented, and non-reproducible. ',
    'She does not identify which third-party tools she used, does not provide crawl parameters, does not document her deduplication methods, provides no downloadable reports as evidence, and admits her January 2026 analysis is only a \u201cqualitative comparison.\u201d Her own report identifies the methodological standards she failed to meet.')

add_numbered_item(doc,
    'Brown repeatedly conflates organic search engine behavior with intentional misconduct. ',
    'That Google\u2019s algorithm may surface DunhamLaw.com in response to a \u201cDunnam\u201d query is a function of algorithmic name-similarity matching, not evidence of targeting. Defendants receive zero clicks from \u201cDunnam\u201d-intent searches. Moreover, even if Defendants were bidding on \u201cDunnam\u201d as a paid keyword\u2014which they are not\u2014competitor brand bidding is a common and accepted practice explicitly permitted by Google, and would not itself violate the 1998 Judgment.')

add_numbered_item(doc,
    'Brown\u2019s corrective advertising estimates ($65,000\u2013$120,000 per year) are unsupported, inflated, and premature. ',
    'No methodology, market data, or comparable cases are cited. The figures are grossly disproportionate to a regional Texas market and presume proven violations that have not been established.')

add_numbered_item(doc,
    'Defendants have taken affirmative steps beyond what the 1998 Judgment requires. ',
    'Defendants voluntarily redirect DunhamLaw.com to DunhamJones.com for users in restricted counties. A Screaming Frog crawl confirms that 66 of 71 restricted-county URLs on DunhamLaw.com redirect via 301 to DunhamJones.com\u2014a systematic, proactive compliance measure that no court order compels.')

add_numbered_item(doc,
    'Brown\u2019s repeated demands for Defendants\u2019 internal documentation appear to be a fishing expedition. ',
    'The items she claims require discovery\u2014metadata, redirects, paid ad presence, county targeting\u2014are all publicly verifiable through the same tools she claims to use. Her insistence on obtaining internal Google Ads account exports, SEO vendor communications, Search Console data, and full site backups goes well beyond what is necessary to assess compliance and would provide Plaintiff\u2019s team with valuable competitive intelligence about Defendants\u2019 marketing strategy.')

# ============================================================
# III. INDEPENDENT ANALYSIS
# ============================================================

doc.add_page_break()
doc.add_heading('III. Independent Analysis: Tools, Methods, and Key Findings', level=1)

add_para(doc, 'Unlike the Brown Report, which names no specific tools and provides no raw data, I have conducted independent analysis using named, industry-standard tools with full documentation. The following summarizes the methods and results.')

doc.add_heading('A. Screaming Frog SEO Spider \u2014 Full Site Crawl of DunhamLaw.com', level=2)

p = doc.add_paragraph()
run = p.add_run('Date: ')
run.bold = True
p.add_run('March 1, 2026  |  ')
run = p.add_run('Duration: ')
run.bold = True
p.add_run('40 minutes  |  ')
run = p.add_run('Scope: ')
run.bold = True
p.add_run('Full site crawl')

add_para(doc, 'Screaming Frog is a widely recognized website auditing tool used by SEO professionals to crawl websites and extract on-page data including page titles, meta descriptions, headings, body text, metadata, and custom search patterns. I configured the crawler to perform a custom search extraction specifically targeting the keyword \u201cdunnam\u201d across three content layers: raw HTML source code, rendered page text, and content area text.')

add_key_finding(doc, 'Across all 1,216 indexable HTML pages on DunhamLaw.com, the custom search extraction returned zero matches for the keyword \u201cdunnam\u201d in HTML, page text, and content areas. The word \u201cdunnam\u201d does not appear anywhere on the DunhamLaw.com website\u2014not in page content, not in metadata, not in source code, not in headers, not in footers, not in navigation elements.')

add_para(doc, 'This finding directly contradicts Brown\u2019s claims of \u201ccontinued algorithmic association with the Dunnam & Dunnam brand\u201d through Defendants\u2019 website. If Defendants were attempting to manipulate search rankings for the \u201cDunnam\u201d brand\u2014as Brown implies\u2014they would need to include that term somewhere in their website content, metadata, or structure. The complete absence of this term, verified through a comprehensive automated crawl, demonstrates that no such manipulation exists.')

add_para(doc, 'The crawl also revealed the redirect structure for restricted-county content:')

add_table(doc,
    ['Category', 'Total URLs', 'Live (200)', 'Redirect to DunhamJones.com (301)'],
    [
        ['Restricted county pages (all 16 counties)', '71', '5*', '66'],
        ['Waco law office page', '1', '0', '1'],
        ['Waco criminal attorneys page', '1', '0', '1'],
    ])

p = doc.add_paragraph()
run = p.add_run('* Of the 5 live URLs: 3 are image files from a November 2025 \u201cNot Guilty\u201d verdict announcement, 1 is a case result page, and 1 is a Tennessee (not Texas) county page. No live, indexable service pages targeting restricted Texas counties remain on DunhamLaw.com.')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
p.paragraph_format.space_after = Pt(8)

add_para(doc, 'Importantly, DunhamLaw.com contains warrant, bail bond, and injury lawyer pages for all 254 Texas counties\u2014not just the 16 restricted counties. The restricted-county pages are part of a blanket statewide coverage strategy, and the vast majority have been redirected to DunhamJones.com as a compliance measure.')

doc.add_heading('B. Google Ads Transparency Center \u2014 Paid Ad Verification', level=2)

add_para(doc, 'The Google Ads Transparency Center is a publicly accessible tool launched by Google that allows any user to view all advertisements run by any verified advertiser. I reviewed all advertisements associated with the domain DunhamLaw.com.')

add_key_finding(doc, 'Every advertisement displayed in the Google Ads Transparency Center for the DunhamLaw.com domain identifies the advertiser as \u201cDUNHAM & JONES ATTORNEYS AT LAW, P.C.\u201d\u2014a verified advertiser. No advertisements use the name \u201cDunham Law Firm,\u201d \u201cDunham, Attorney at Law,\u201d \u201cDunhams,\u201d or any other prohibited name formulation. No advertisements reference the \u201cDunnam\u201d brand in any way.')

add_para(doc, 'Brown could have\u2014and should have\u2014consulted this tool before opining on Defendants\u2019 paid advertising. It is free, publicly accessible, and specifically designed to provide transparency into advertiser identity. Her failure to reference it is a significant methodological omission.')

doc.add_heading('C. SimilarWeb \u2014 Paid and Organic Keyword Analysis', level=2)

add_para(doc, 'SimilarWeb is a leading competitive intelligence platform used to analyze website traffic, keyword rankings, and advertising activity. I analyzed paid and organic keyword data for DunhamLaw.com, DunhamJones.com, and DunnamLaw.com covering February 2025 through January 2026.')

add_key_finding(doc, 'Neither DunhamLaw.com nor DunhamJones.com targeted any paid keywords containing \u201cdunnam\u201d or any prohibited name formulation during the 12-month period analyzed.')

add_key_finding(doc, 'DunhamLaw.com appeared organically for 3 queries containing \u201cdunnam\u201d\u2014all with zero clicks. DunhamJones.com appeared for 1 \u201cdunnam\u201d query\u2014also with zero clicks. These incidental appearances generate no traffic and are the natural result of Google\u2019s fuzzy matching of similar names.')

add_para(doc, 'It is important to note that even if Defendants were bidding on \u201cDunnam & Dunnam\u201d as a paid keyword\u2014which the data confirms they are not\u2014this would still not constitute a violation of the 1998 Judgment. The Judgment prohibits using certain name formulations in advertising, not bidding on a competitor\u2019s brand name in a search advertising platform. Competitor brand bidding is a common and accepted practice in paid search advertising. Google explicitly permits advertisers to bid on competitors\u2019 brand names as keywords; it is a standard feature of the Google Ads platform used across virtually every industry. The fact that Defendants have not engaged in this entirely permissible practice further underscores the absence of any intent to trade on the Dunnam & Dunnam brand.')

add_para(doc, 'If there were any evidence of Defendants attempting to conflate their brand with Dunnam & Dunnam, the clearest and most direct proof would be paid search bids targeting the \u201cDunnam\u201d brand\u2014purchasing the right to appear in search results when users search for \u201cDunnam & Dunnam.\u201d No such bids exist. This is not a gap in the evidence; it is affirmative proof that no such targeting has occurred.')

doc.add_heading('D. SEMrush \u2014 Organic Positions and Backlink Analysis', level=2)

add_para(doc, 'SEMrush is one of the world\u2019s most widely used SEO research platforms, trusted by enterprises and agencies globally. I analyzed organic ranking positions for all three domains and performed a comprehensive backlink and anchor text analysis for DunhamLaw.com.')

p = doc.add_paragraph()
run = p.add_run('Organic Rankings for \u201cDunnam\u201d Keywords:')
run.bold = True
p.paragraph_format.space_after = Pt(4)

add_table(doc,
    ['Domain', '"Dunnam" Keywords', 'Positions', 'Traffic from "Dunnam" Keywords'],
    [
        ['dunhamlaw.com', '3', '19, 79', '0'],
        ['dunhamjones.com', '12', '21\u201373 (all page 3+)', '0'],
    ])

add_para(doc, 'All \u201cdunnam\u201d keyword rankings for Defendants\u2019 domains fall on page 3 or deeper of Google search results (position 21+). These positions are functionally invisible to searchers\u2014fewer than 1% of users click past page 2\u2014and they generate zero measurable traffic. These are passive, incidental rankings caused by algorithmic name similarity, not targeted optimization.')

add_key_finding(doc, 'Across 28,279 total backlinks pointing to DunhamLaw.com (1,807 unique anchor texts), zero contain the word \u201cdunnam\u201d in any form. Not a single backlink uses \u201cdunnam,\u201d \u201cdunnam & dunnam,\u201d \u201cdunnam and dunnam,\u201d or any variation of the Dunnam brand as anchor text.')

add_para(doc, 'If Defendants were engaged in any link-building strategy to associate their domain with the \u201cDunnam\u201d brand\u2014a common SEO tactic for brand confusion\u2014this would be plainly visible in the anchor text data. The complete absence of \u201cdunnam\u201d from 28,279 backlinks is definitive evidence that no such strategy exists or has ever existed.')

add_para(doc, 'All \u201cDunham\u201d-containing anchor texts reference Defendants\u2019 legitimate brand names: \u201cDunham & Jones,\u201d \u201cDunham & Jones Attorneys at Law, P.C.,\u201d \u201cPaul Dunham,\u201d and \u201cDunhamLaw.com.\u201d')

# ============================================================
# IV. MISCHARACTERIZATION
# ============================================================

doc.add_page_break()
doc.add_heading('IV. The Brown Report Mischaracterizes the 1998 Judgment', level=1)

add_para(doc, 'A proper rebuttal must begin with what the 1998 Agreed Judgment actually says, because the Brown Report consistently extends its reach beyond the four corners of that order.')

doc.add_heading('A. What the 1998 Judgment Prohibits', level=2)

add_para(doc, 'The 1998 Judgment permanently enjoins Defendants from using, directly or indirectly, in 16 specified Texas counties (Bell, McLennan, Coryell, Hill, Bosque, Johnson, Brazos, Hamilton, Somervell, Limestone, Freestone, Robertson, Falls, Madison, Navarro, and Leon), the following names: \u201cDunham, Attorney at Law,\u201d \u201cDunhams, Attorney at Law,\u201d \u201cDunham Law Firm,\u201d \u201cDunham Firm,\u201d \u201cDunhams,\u201d \u201cDunham and Associates,\u201d and variations thereof using the spelling \u201cDunnam\u201d rather than \u201cDunham.\u201d It further prohibits advertising or telephone listings using these specified names in the specified area.')

doc.add_heading('B. What the 1998 Judgment Does Not Prohibit', level=2)

add_bullet(doc, ' \u201cDunham\u201d is the Defendants\u2019 legal surname. The Judgment restricts specific name formulations, not the surname in isolation. The domain DunhamLaw.com is not one of the prohibited names.', 'The use of the surname \u201cDunham\u201d itself.')

add_bullet(doc, ' Defendants are free to practice law in all 16 counties; they are restricted only in the names under which they advertise.', 'The practice of law in the specified counties.')

add_bullet(doc, ' The name \u201cDunham & Jones\u201d is not among the prohibited names. Defendants may advertise their services in the specified counties under this name. As confirmed by the Google Ads Transparency Center, all of Defendants\u2019 advertisements identify them as \u201cDunham & Jones Attorneys at Law, P.C.\u201d', 'Advertising under non-prohibited names.')

add_bullet(doc, ' The 1998 Judgment was entered years before modern search engines existed in their current form. It addresses naming practices and advertising\u2014not how third-party algorithms index or display information about Defendants.', 'Organic search engine rankings.')

add_bullet(doc, ' Dunham & Jones operates offices and serves clients in the restricted counties. It is standard and necessary practice for a law firm to mention the locations where it operates on its website. The Judgment prohibits using certain names in those areas\u2014not mentioning the areas themselves.', 'Mentioning geographic locations on their website.')

add_bullet(doc, ' The Judgment does not require Defendants to suppress their legal name from search engine indexes or to implement negative keywords against their own surname.', 'The existence of SEO signals, metadata, or domain architecture.')

doc.add_heading('C. Brown\u2019s Overreach', level=2)

add_para(doc, 'Brown\u2019s report effectively asks this Court to expand the 1998 Judgment to prohibit: the operation of the domain DunhamLaw.com, organic search visibility for the surname \u201cDunham\u201d in restricted counties, any mention of restricted counties on Defendants\u2019 website, and any \u201calgorithmic association\u201d between Defendants and Plaintiff. None of these requirements appear in the 1998 Judgment, and imposing them would effectively prohibit Defendants from operating a law practice under their own legal name.')

add_para(doc, 'Brown\u2019s keyword frequency counts\u2014\u201c450 mentions of \u2018Waco criminal attorney,\u2019 195 mentions of \u2018Bell County\u2019\u201d\u2014are presented as though they constitute violations. They do not. Having content about locations where Defendants practice law is not a violation of the Judgment and is not connected to the prohibited naming practices. Dunham & Jones is allowed to operate in those locations, and it is only logical that they would reference them on their website.')

# ============================================================
# V. EVIDENCE OF NON-VIOLATION
# ============================================================

doc.add_page_break()
doc.add_heading('V. The Evidence Confirms No Violation Exists', level=1)

doc.add_heading('A. No Prohibited Names on Defendants\u2019 Websites', level=2)

add_para(doc, 'The Screaming Frog crawl of DunhamLaw.com\u20142,104 internal URLs, 1,548 HTML pages, custom extraction for \u201cdunnam\u201d across HTML source, page text, and content areas\u2014returned zero matches. The word \u201cdunnam\u201d does not appear anywhere on the site.')

add_para(doc, 'This is the definitive test. If Defendants were using prohibited names on their website, were embedding \u201cDunnam\u201d in metadata, or were targeting the \u201cDunnam & Dunnam\u201d brand through any on-page element, a comprehensive site crawl would detect it. This analysis was available to Brown using a tool she claims to use\u2014Screaming Frog is listed on her CV\u2014yet she either did not perform this crawl or chose not to report the results.')

doc.add_heading('B. No Prohibited Advertisements', level=2)

add_para(doc, 'The Google Ads Transparency Center confirms that every paid advertisement associated with the DunhamLaw.com domain identifies the advertiser as \u201cDunham & Jones Attorneys at Law, P.C.\u201d\u2014a name that is not prohibited by the 1998 Judgment. No advertisements use \u201cDunham Law Firm,\u201d \u201cDunham, Attorney at Law,\u201d \u201cDunhams,\u201d or any other prohibited formulation. No advertisements reference \u201cDunnam\u201d in any way.')

add_para(doc, 'SimilarWeb data further confirms that neither DunhamLaw.com nor DunhamJones.com targeted any paid keywords containing \u201cdunnam\u201d or any prohibited name formulation during the 12-month period from February 2025 through January 2026.')

add_para(doc, 'Brown concedes she \u201ccannot verify\u201d paid advertising activity but still implies violation through \u201cbrand-adjacent visibility\u201d language. If Defendants were running paid search ads using prohibited names in restricted counties, this would be visible to anyone\u2014including Brown herself\u2014through a simple Google search from within those counties, through the Google Ads Transparency Center, or through any third-party competitive intelligence tool. The fact that no such evidence exists is not a gap in the record; it is evidence that no violation has occurred.')

doc.add_heading('C. No \u201cDunnam\u201d Brand Targeting in SEO', level=2)

add_para(doc, 'The SEMrush backlink analysis provides additional confirmation. Across 28,279 backlinks to DunhamLaw.com, zero use the word \u201cdunnam\u201d in anchor text. This is significant because anchor text manipulation\u2014building links with competitor brand names as anchor text\u2014is the primary SEO technique for creating brand association. The complete absence of \u201cdunnam\u201d from Defendants\u2019 entire backlink profile demonstrates that no such strategy has been employed.')

add_para(doc, 'All three lines of evidence\u2014on-page content (Screaming Frog), paid advertising (Google Ads Transparency Center, SimilarWeb), and off-page SEO (SEMrush backlinks)\u2014independently confirm the same conclusion: Defendants have not targeted, referenced, or attempted to associate with the \u201cDunnam\u201d brand.')

doc.add_heading('D. The \u201cIf It\u2019s Visible to Google, It\u2019s Visible to You\u201d Principle', level=2)

add_para(doc, 'A foundational principle of search engine optimization is that search engines can only index and rank content that is publicly accessible. If Defendants were using prohibited names on their website, embedding them in metadata, or targeting the \u201cDunnam & Dunnam\u201d brand through on-page content, this information would be visible to any person using a web browser, any search engine crawler, and any third-party SEO tool. The fact that Brown cannot produce hard evidence of any specific violation\u2014despite having access to the same public internet as everyone else\u2014strongly suggests that no such violation exists.')

# ============================================================
# VI. METHODOLOGY
# ============================================================

doc.add_page_break()
doc.add_heading('VI. Brown\u2019s Methodology Is Insufficient and Non-Reproducible', level=1)

doc.add_heading('A. Unnamed Tools and Undocumented Processes', level=2)

add_para(doc, 'Brown references \u201cthird-party SEO visibility data\u201d and \u201cthird-party keyword visibility\u201d throughout her report without ever identifying the specific tools used. Industry-standard practice for expert analysis requires naming tools (e.g., Ahrefs, SEMrush, SimilarWeb, Moz), specifying the exact date of data extraction, and providing the raw report outputs. All of these tools generate downloadable reports that serve as verifiable evidence. Brown provides none.')

add_para(doc, 'By contrast, every finding in this report identifies the specific tool used, the date of analysis, the scope of data collected, and the availability of raw exports for verification.')

doc.add_heading('B. Self-Identified Methodological Failures', level=2)

add_para(doc, 'In Section XI of her report, Brown herself acknowledges that proper reproducibility would require documenting:')

add_bullet(doc, 'Crawl scope and user-agent')
add_bullet(doc, 'Date range')
add_bullet(doc, 'Deduplication method for templated content')
add_bullet(doc, 'Handling of dynamic and paginated content')
add_bullet(doc, 'Phrase matching and variant handling logic')

add_para(doc, 'Brown admits she did not provide any of these parameters. This is a significant admission. Without a documented deduplication methodology, her keyword frequency counts (e.g., \u201c450 mentions of \u2018Waco criminal attorney\u2019\u201d) are almost certainly inflated by templated content\u2014headers, footers, navigation menus, and sidebar elements that repeat across every page of a website. A single navigation element containing \u201cWaco\u201d that appears on 450 pages would be counted as \u201c450 mentions\u201d under an undeduplicated methodology, despite representing a single design choice.')

doc.add_heading('C. Qualitative Conclusions from Quantitative Claims', level=2)

add_para(doc, 'Brown admits that her January 2026 update is only a \u201cqualitative comparison\u201d to the August 2025 baseline, yet she draws the definitive conclusion that there has been \u201cno material retreat\u201d in restricted-area visibility. An expert cannot make quantitative claims (\u201cno material retreat\u201d) based on admittedly qualitative analysis. Her approximate keyword counts, which she characterizes as \u201cdescriptive indicators\u201d that \u201cshould not be treated as a definitive census,\u201d cannot support the categorical conclusions she draws elsewhere in the report.')

doc.add_heading('D. Publicly Verifiable Items Treated as Requiring Discovery', level=2)

add_para(doc, 'Brown\u2019s Verification/Proof Matrix (Attachment A) claims that several items are \u201cnot publicly verifiable\u201d and require production from Defendants. In fact, many of these items are partially or fully verifiable through public tools:')

add_bullet(doc, ' Fully verifiable through a Screaming Frog crawl\u2014a tool Brown claims to use. I performed this crawl and found zero instances of \u201cdunnam.\u201d', 'Metadata cleanup:')
add_bullet(doc, ' Verifiable through direct browser testing, HTTP header inspection, and redirect-checking tools. My crawl confirmed 66 of 71 restricted-county URLs redirect to DunhamJones.com.', 'Domain redirects:')
add_bullet(doc, ' Verifiable through the Google Ads Transparency Center, direct search observation in restricted counties, and third-party paid search intelligence tools. My analysis of the Transparency Center confirmed all ads use \u201cDunham & Jones Attorneys at Law, P.C.\u201d', 'Paid ad presence:')
add_bullet(doc, ' Verifiable through site crawling and content analysis. My Screaming Frog crawl captured every page and its content.', 'County targeting on-site:')

add_para(doc, 'That Brown claims these items require discovery from Defendants, rather than conducting the public analysis herself, raises serious questions about the thoroughness of her investigation. Either she lacks familiarity with the tools that would answer these questions, or she recognizes that the publicly available evidence does not support her conclusions. In either case, her request for extensive internal documentation\u2014including full site backups, SEO vendor communications, Google Analytics data, Search Console data, and complete Google Ads account exports\u2014goes well beyond what is necessary to assess compliance with a naming restriction. Such production would effectively provide Plaintiff\u2019s team with a comprehensive roadmap of Defendants\u2019 marketing strategy and competitive intelligence, which has nothing to do with the 1998 Judgment\u2019s narrow prohibition on specific name formulations.')

# ============================================================
# VII. ORGANIC RANKINGS
# ============================================================

doc.add_page_break()
doc.add_heading('VII. Organic Search Rankings Are Not Evidence of Misconduct', level=1)

doc.add_heading('A. How Search Engines Handle Similar Names', level=2)

add_para(doc, 'Brown\u2019s finding that DunhamLaw.com ranked for queries like \u201cdunnam & dunnam\u201d is presented as evidence of wrongdoing. It is not. Search engines, including Google, routinely surface results for approximate name matches, misspellings, and phonetically similar terms. When a user searches for \u201cdunnam,\u201d Google may surface results for \u201cdunham\u201d because the algorithm recognizes the names as similar. This is core algorithmic behavior that neither party controls.')

add_para(doc, 'The closest Defendants could come to influencing rankings for the \u201cDunnam & Dunnam\u201d brand would be to include that name in their website content, metadata, or link structure. My Screaming Frog crawl confirms this has never been the case: zero instances of \u201cdunnam\u201d in 1,548 HTML pages. My SEMrush analysis confirms zero instances of \u201cdunnam\u201d in 28,279 backlinks.')

doc.add_heading('B. Neither Party Controls Google\u2019s Algorithm', level=2)

add_para(doc, 'It is a fundamental reality of modern search that no website operator controls how Google ranks their site for queries they have not specifically targeted. Google\u2019s ranking algorithm considers hundreds of factors, including domain name similarity, geographic proximity, practice area overlap, and user behavior signals. That two law firms with near-identical surnames, operating in overlapping geographic markets and the same practice areas, would appear in each other\u2019s branded search results is entirely expected and does not indicate any manipulation by either party.')

doc.add_heading('C. The \u201cAlgorithmic Ambiguity\u201d Argument Is Without Merit', level=2)

add_para(doc, 'Brown introduces the concept of \u201calgorithmic ambiguity\u201d and describes a chatbot incident where \u201cincorrect or conflicting contact information\u201d was surfaced. She admits this was \u201cnot consistently reproducible\u201d and is \u201cnot relied upon as a standalone finding.\u201d This admission effectively concedes the finding has no evidentiary value. Defendants do not control third-party AI chatbot behavior, and isolated, non-reproducible algorithmic outputs are not evidence of misconduct.')

# ============================================================
# VIII. PROACTIVE COMPLIANCE
# ============================================================

doc.add_heading('VIII. Defendants Have Taken Proactive Steps Beyond the Judgment\u2019s Requirements', level=1)

add_para(doc, 'Brown\u2019s report fails to give adequate weight to a significant fact: Defendants voluntarily redirect their primary domain DunhamLaw.com to DunhamJones.com for users located in the restricted counties.')

add_para(doc, 'My Screaming Frog crawl independently confirms the scope of this redirect program:')

add_bullet(doc, ' on DunhamLaw.com redirect via HTTP 301 to DunhamJones.com', '66 of 71 restricted-county URLs')
add_bullet(doc, 'Redirects cover all 16 restricted counties across three content categories: warrants, bail bonds, and injury lawyers')
add_bullet(doc, ' redirects to DunhamJones.com', 'The Waco law office page (/law-offices/waco-law-office/)')
add_bullet(doc, ' redirects to DunhamJones.com', 'The Waco criminal attorneys page (/tx/waco-criminal-attorneys/)')
add_bullet(doc, 'City-specific pages for Bryan, Harker Heights, Copperas Cove, Hillsboro, Cleburne, and Marlin all redirect to DunhamJones.com')

add_para(doc, 'This is an extraordinary measure. Defendants operate a successful law practice with a 20+ year established domain. No provision of the 1998 Judgment requires them to redirect their primary domain for any geographic subset of users. That they have chosen to do so\u2014at the cost of diminished brand consistency and potential user confusion\u2014demonstrates a good-faith commitment to compliance that goes well beyond the Judgment\u2019s requirements. In no other context would a successful business voluntarily redirect its primary, long-established domain for a subset of its service areas.')

add_para(doc, 'In my professional experience managing enterprise-level web properties, geographic domain redirection is a technically sophisticated solution that requires ongoing maintenance and investment. It is the opposite of the \u201cselective implementation choices\u201d that Brown alleges.')

# ============================================================
# IX. CORRECTIVE ADVERTISING
# ============================================================

doc.add_page_break()
doc.add_heading('IX. Brown\u2019s Corrective Advertising Estimates Are Unsupported and Inflated', level=1)

doc.add_heading('A. No Methodology or Basis Provided', level=2)

add_para(doc, 'Brown proposes annual corrective advertising costs of $65,000 to $120,000 per year for an unspecified multi-year period, without citing any methodology, market data, rate cards, comparable cases, or industry benchmarks. The figures appear to be arbitrary.')

doc.add_heading('B. Disproportionate to the Market', level=2)

add_para(doc, 'Waco, Texas (McLennan County) is a regional market with a metropolitan population of approximately 270,000. The proposed \u201cpaid search corrective ads\u201d budget of $25,000\u2013$50,000 per year is orders of magnitude beyond what would be required for a legitimate search clarification campaign in a market of this size. For context, $50,000 in annual Google Ads spend in a regional legal market would purchase an extraordinarily dominant advertising presence\u2014far beyond any reasonable \u201ccorrective\u201d purpose.')

doc.add_heading('C. Self-Serving Monitoring Fees', level=2)

add_para(doc, 'Brown proposes $12,000\u2013$20,000 per year for \u201congoing monitoring and enforcement.\u201d At her stated rate of $125 per hour, this represents 96 to 160 hours of annual monitoring\u2014roughly 2 to 3 hours per week, every week, indefinitely. This is grossly disproportionate to the monitoring needs of a single regional trademark matter and appears designed to create a recurring revenue stream rather than address a genuine compliance need.')

doc.add_heading('D. Search Signals Do Not Require Multi-Year Correction', level=2)

add_para(doc, 'Brown asserts that search engine associations \u201cdo not self-correct simply because conduct stops\u201d and that corrective efforts \u201cmust be sustained\u201d over a \u201cmulti-year period.\u201d This is misleading. Search engine results and ranking signals are highly dynamic. Google recrawls and reindexes active websites on a continuous basis, typically within days to weeks. Historical ranking signals naturally decay as content is updated, removed, or superseded. A multi-year corrective campaign is not supported by how search engines actually function.')

doc.add_heading('E. Premature and Presumptive', level=2)

add_para(doc, 'The entire corrective advertising analysis presumes that violations have been proven and damages have been established. Neither is the case. Calculating corrective advertising costs before establishing that any violation occurred puts the cart before the horse and should be disregarded.')

# ============================================================
# X. QUALIFICATIONS
# ============================================================

doc.add_heading('X. Observations on the Opposing Expert\u2019s Qualifications', level=1)

add_para(doc, 'While the weight to be given expert testimony is a matter for the Court, I note the following regarding Brown\u2019s qualifications as they relate to the reliability of her opinions:')

add_bullet(doc, ' not search marketing or advertising expertise.', 'Brown\u2019s professional background prior to founding Digital Media Butterfly in 2012 was in system administration\u2014managing phone systems, hardware rollouts, and providing technical support for approximately 400 users. This is information technology work,')

add_bullet(doc, ' in web design and development. She holds no degree in marketing, advertising, computer science, or any related field.', 'Brown\u2019s formal education consists of coursework in a 1998 night program')

add_bullet(doc, ' in the field of search marketing, SEO, or digital advertising.', 'Brown has never previously testified as an expert witness and has no publications')

add_bullet(doc, '\u2014no enterprise, Fortune 500, or other significant engagements that would demonstrate experience with the kind of multi-domain, multi-market analysis she attempts in this report.', 'Brown names no clients in her CV')

add_bullet(doc, ' that exceed the scope of technical expertise and invade the province of the fact-finder. An expert in search marketing can opine on technical facts; she cannot opine on whether Defendants\u2019 \u201cselective implementation choices\u201d constitute intentional non-compliance.', 'Brown\u2019s report makes legal conclusions about intent, willfulness, and compliance interpretation')

add_bullet(doc, '\u2014the same city as Plaintiff Dunnam & Dunnam\u2014which, while not dispositive, is relevant to evaluating independence in a local trademark dispute.', 'Brown is located in Waco, Texas')

add_bullet(doc, '\u2014such as a site-wide content crawl for the \u201cdunnam\u201d keyword or a backlink anchor text analysis\u2014that would have directly addressed her central allegations. This raises questions about her proficiency with these tools and the thoroughness of her investigation.', 'Brown claims to use tools such as Screaming Frog and Ahrefs, yet did not perform the basic analyses these tools enable')

# ============================================================
# XI. CONCLUSION
# ============================================================

doc.add_page_break()
doc.add_heading('XI. Conclusion', level=1)

add_para(doc, 'The Brown Report presents 15 pages of speculation, unsupported claims, and methodological shortcuts in an attempt to create the appearance of non-compliance where none has been demonstrated. My independent analysis using industry-standard tools has produced verifiable evidence that directly contradicts Brown\u2019s central allegations:')

add_bullet(doc, ' A Screaming Frog crawl of all 1,216 indexable pages found zero instances of \u201cdunnam\u201d in HTML, page text, or content areas.', 'No \u201cDunnam\u201d content on Defendants\u2019 website.')

add_bullet(doc, ' The Google Ads Transparency Center confirms all advertisements use \u201cDunham & Jones Attorneys at Law, P.C.\u201d\u2014a non-prohibited name. SimilarWeb confirms no paid keywords targeting \u201cdunnam\u201d or any prohibited name on either domain.', 'No prohibited advertising.')

add_bullet(doc, ' SEMrush analysis of 28,279 backlinks found zero anchor texts containing \u201cdunnam.\u201d All link signals reference Defendants\u2019 legitimate brand names.', 'No brand targeting through SEO.')

add_bullet(doc, ' Both SimilarWeb and SEMrush confirm that Defendants\u2019 websites receive zero click traffic from \u201cdunnam\u201d-intent keywords. No paid keyword bids target the \u201cDunnam\u201d brand\u2014a practice that, even if it occurred, would be permitted by Google and would not violate the 1998 Judgment.', 'Zero traffic from \u201cDunnam\u201d searches.')

add_bullet(doc, ' 66 of 71 restricted-county URLs on DunhamLaw.com redirect to DunhamJones.com via HTTP 301\u2014a voluntary, proactive compliance measure beyond the Judgment\u2019s requirements.', 'Systematic compliance measures in place.')

add_bullet(doc, ' Brown treats the Judgment as a blanket prohibition on the use of the name \u201cDunham\u201d in digital environments. It is not. The Judgment prohibits specific name formulations in specific counties\u2014a far narrower restriction than Brown applies.', 'Misreading of the 1998 Judgment.')

add_bullet(doc, ' Brown\u2019s analytical framework is structured so that no evidence could ever demonstrate compliance: if she finds something, it\u2019s a violation; if she finds nothing, it requires discovery to confirm. This is not rigorous analysis\u2014it is advocacy.', 'Unfalsifiable reasoning.')

add_bullet(doc, ' The corrective advertising estimates of $65,000\u2013$120,000 per year are unsupported by any methodology, wildly disproportionate to the regional market, and premature absent any established violation.', 'Inflated damages.')

add_para(doc, 'In my professional opinion, Defendants\u2019 online properties do not demonstrate the violations alleged in the Brown Report. The evidence demonstrates the opposite: Defendants have systematically avoided using prohibited names on their website, in their advertising, and in their SEO practices, and they have implemented voluntary compliance measures that go beyond the 1998 Judgment\u2019s requirements. The Brown Report should be given little weight due to its methodological deficiencies, unsupported conclusions, and fundamental mischaracterization of both the applicable legal standard and the technical realities of search engine behavior.')

# ============================================================
# SIGNATURE
# ============================================================

add_para(doc, '')
add_para(doc, 'The opinions expressed herein are based on my professional experience, analysis of publicly available information, and independent research using industry-standard tools as described herein. Raw tool exports supporting the findings in this report are available upon request. I reserve the right to supplement this report upon receipt of additional materials or discovery.')

add_para(doc, '')
add_para(doc, 'Executed on __________________, in Los Angeles, California.')

add_para(doc, '')
add_para(doc, '')
add_para(doc, '________________________________________')

p = doc.add_paragraph()
run = p.add_run('Kenny Hyder')
run.bold = True
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph('Founder, Hyder Media')
p.paragraph_format.space_after = Pt(0)
p = doc.add_paragraph('kenny@hyder.me | (619) 850-7217')
p.paragraph_format.space_after = Pt(0)
p = doc.add_paragraph('hyder.me')

# Save
output = os.path.join(os.path.dirname(__file__), 'expert-rebuttal-report-v2.docx')
doc.save(output)
print(f'Word document saved to {output}')
