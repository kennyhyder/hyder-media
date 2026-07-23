#!/usr/bin/env python3
"""Generate Expert Rebuttal Report as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading 1 style (section headers)
h1 = doc.styles['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Pt(14)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(8)
h1.paragraph_format.keep_with_next = True
# Remove numbering
h1_element = h1.element
rPr = h1_element.find(qn('w:rPr'))

# Heading 2 style (subsection headers)
h2 = doc.styles['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Pt(12)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
h2.paragraph_format.space_before = Pt(16)
h2.paragraph_format.space_after = Pt(6)
h2.paragraph_format.keep_with_next = True


def add_bold_text(paragraph, bold_text, normal_text=""):
    run = paragraph.add_run(bold_text)
    run.bold = True
    if normal_text:
        paragraph.add_run(normal_text)


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p


# ============================================================
# HEADER / TITLE PAGE
# ============================================================

# Logo
logo_path = os.path.join(os.path.dirname(__file__), '..', 'assets', 'imgs', 'logos', 'hyder-media-logo.png')
if os.path.exists(logo_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(0.6))

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Expert Rebuttal Report')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Response to Expert Report of Caryn Brown, Digital Media Butterfly\ndated February 6, 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)

# Separator
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('_' * 60)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

# Meta info
meta_items = [
    ("Prepared for: ", "Dunham Law Firm, P.C. / Dunham & Jones L.L.P."),
    ("Prepared by: ", "Kenny Hyder, Founder, Hyder Media"),
    ("Case: ", "Dunnam & Dunnam LLP v. Dunham Law Firm, P.C., et al."),
    ("", "Civil Action No. 6:21-cv-1041-ADA-DGT, W.D. Texas, Waco Division"),
    ("Date: ", "[DATE]"),
]
for bold_part, normal_part in meta_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    if bold_part:
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(10)
    run = p.add_run(normal_part)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)

doc.add_page_break()

# ============================================================
# I. INTRODUCTION
# ============================================================

doc.add_heading('I. Introduction and Qualifications', level=1)

add_para(doc, 'I have been retained by counsel for Defendants to review and respond to the Expert Report of Caryn Brown ("Brown Report"), dated February 6, 2026, prepared on behalf of Plaintiff Dunnam & Dunnam LLP. My analysis addresses the technical accuracy, methodological rigor, and factual basis of Brown\'s claims regarding Defendants\' online properties and their relationship to the Agreed Judgment dated April 24, 1998 (the "1998 Judgment").')

add_para(doc, 'I am a digital marketing and search advertising consultant with over 20 years of hands-on experience in search engine optimization (SEO), paid search advertising (Google Ads/PPC), content marketing, online reputation management, and competitive analysis. I have directed search marketing strategy for multiple Fortune 500 companies, including Fortune 20 and Fortune 5 businesses. I am a recognized contributing author at Search Engine Journal, a veteran speaker at major industry conferences including Pubcon, SMX, Affiliate Summit, and the Direct Marketing Association, and I serve as a judge for the US Search Awards. My full curriculum vitae is attached.')

add_para(doc, 'My opinions are based on my professional experience, direct observation using industry-standard tools and methodologies, and review of the Brown Report, the 1998 Agreed Judgment, and the publicly observable online properties at issue. I am being compensated at a rate of $300 per hour for my work in this matter. My compensation is not contingent on the outcome of this case or the opinions expressed herein.')

# ============================================================
# II. SUMMARY
# ============================================================

doc.add_heading('II. Summary of Opinions', level=1)

add_para(doc, 'After thorough review of the Brown Report, the underlying court orders, and independent analysis of Defendants\' online properties, it is my opinion that:')

add_bullet(doc, "The Brown Report fails to identify a single concrete violation of the 1998 Judgment. ", "")
p = doc.paragraphs[-1]
p.clear()
add_bold_text(p, "The Brown Report fails to identify a single concrete violation of the 1998 Judgment. ", "Despite 15 pages of analysis, Brown does not produce a screenshot, an advertisement, a search result, or any verifiable artifact demonstrating that Defendants have used the prohibited names in the prohibited counties. If such evidence existed, it would be trivially easy to capture and present.")

add_bullet(doc, "", "")
p = doc.paragraphs[-1]
p.clear()
add_bold_text(p, "The Brown Report fundamentally mischaracterizes the scope of the 1998 Judgment. ", 'The Judgment restricts the use of specific name formulations (e.g., "Dunham Law Firm," "Dunham, Attorney at Law") in specific Texas counties. It does not prohibit Defendants from using their legal surname "Dunham," from operating the domain DunhamLaw.com, from practicing law in those counties, or from advertising under non-prohibited names such as "Dunham & Jones."')

add_bullet(doc, "", "")
p = doc.paragraphs[-1]
p.clear()
add_bold_text(p, "Brown's methodology is insufficient, undocumented, and non-reproducible. ", 'She does not identify which third-party tools she used, does not provide crawl parameters, does not document her deduplication methods, and admits her January 2026 analysis is only a "qualitative comparison." Her own report identifies the methodological standards she failed to meet.')

add_bullet(doc, "", "")
p = doc.paragraphs[-1]
p.clear()
add_bold_text(p, "Brown repeatedly conflates organic search engine behavior with intentional misconduct. ", 'That Google\'s algorithm may surface DunhamLaw.com in response to a "Dunnam" query is a function of algorithmic name-similarity matching, not evidence that Defendants have targeted or manipulated the Dunnam & Dunnam brand.')

add_bullet(doc, "", "")
p = doc.paragraphs[-1]
p.clear()
add_bold_text(p, "Brown's corrective advertising estimates ($65,000\u2013$120,000 per year) are unsupported, inflated, and premature. ", "No methodology, market data, or comparable cases are cited. The figures are grossly disproportionate to a regional Texas market and presume proven violations that have not been established.")

add_bullet(doc, "", "")
p = doc.paragraphs[-1]
p.clear()
add_bold_text(p, "Defendants have taken affirmative steps beyond what the 1998 Judgment requires. ", "Notably, Defendants voluntarily redirect their primary domain DunhamLaw.com to DunhamJones.com for users in restricted counties\u2014a significant, proactive measure that no court order compels and that Brown's report fails to give appropriate weight.")

# ============================================================
# III. MISCHARACTERIZATION
# ============================================================

doc.add_page_break()
doc.add_heading('III. The Brown Report Mischaracterizes the 1998 Judgment', level=1)

add_para(doc, "A proper rebuttal must begin with what the 1998 Agreed Judgment actually says, because the Brown Report consistently extends its reach beyond the four corners of that order.")

doc.add_heading('A. What the 1998 Judgment Prohibits', level=2)

add_para(doc, 'The 1998 Judgment permanently enjoins Defendants from using, directly or indirectly, in 16 specified Texas counties, the following names: "Dunham, Attorney at Law," "Dunhams, Attorney at Law," "Dunham Law Firm," "Dunham Firm," "Dunhams," "Dunham and Associates," and variations thereof using the spelling "Dunnam" rather than "Dunham." It further prohibits advertising or telephone listings using these specified names in the specified area.')

doc.add_heading('B. What the 1998 Judgment Does Not Prohibit', level=2)

add_bullet(doc, ' "Dunham" is the Defendants\' legal surname. The Judgment restricts specific name formulations, not the surname in isolation. The domain DunhamLaw.com is not one of the prohibited names.', "The use of the surname \u201cDunham\u201d itself.")

add_bullet(doc, " Defendants are free to practice law in all 16 counties; they are restricted only in the names under which they advertise.", "The practice of law in the specified counties.")

add_bullet(doc, ' The name "Dunham & Jones" is not among the prohibited names. Defendants may advertise their services in the specified counties under this name.', "Advertising under non-prohibited names.")

add_bullet(doc, " The 1998 Judgment was entered years before modern search engines existed in their current form. It addresses naming practices and advertising\u2014not how third-party algorithms index or display information about Defendants.", "Organic search engine rankings.")

add_bullet(doc, " The Judgment does not require Defendants to suppress their legal name from search engine indexes or to implement negative keywords against their own surname.", "The existence of SEO signals, metadata, or domain architecture.")

doc.add_heading('C. Brown\u2019s Overreach', level=2)

add_para(doc, "Brown's report effectively asks this Court to expand the 1998 Judgment to prohibit: the operation of the domain DunhamLaw.com, organic search visibility for the surname \u201cDunham\u201d in restricted counties, any mention of restricted counties on Defendants\u2019 website, and any \u201calgorithmic association\u201d between Defendants and Plaintiff. None of these requirements appear in the 1998 Judgment, and imposing them would effectively prohibit Defendants from operating a law practice under their own legal name.")

# ============================================================
# IV. FAILURE TO PRODUCE EVIDENCE
# ============================================================

doc.add_heading('IV. The Brown Report Fails to Produce Evidence of Any Violation', level=1)

doc.add_heading('A. No Advertisements Captured', level=2)

add_para(doc, "If Defendants were running paid search advertisements using prohibited names in restricted counties, this would be trivially easy to prove. A person physically located in one of the 16 restricted counties could simply search for the prohibited terms on Google and capture a screenshot of any resulting advertisement. Brown does not claim to have done this, and she produces no such screenshot anywhere in her report. The absence of this basic evidence is telling: if such advertisements existed, they would be visible to any member of the public\u2014including Brown herself.")

add_para(doc, "Furthermore, paid search advertisements are now publicly accessible through the Google Ads Transparency Center, which allows any user to view all advertisements run by any verified advertiser. Brown does not reference this tool or provide any data from it. This is a significant omission from someone holding herself out as a search advertising expert.")

doc.add_heading('B. No Prohibited Names Identified on Defendants\u2019 Websites', level=2)

add_para(doc, 'Brown claims that "Dunham" appears in URLs, page titles, headers, and meta content. This is unsurprising\u2014"Dunham" is Defendants\u2019 legal surname and the name of their law firm. However, Brown does not identify a single instance where the prohibited name formulations ("Dunham Law Firm," "Dunham, Attorney at Law," "Dunhams," etc.) appear on Defendants\u2019 websites in the context of advertising in restricted counties.')

add_para(doc, 'This is publicly verifiable. A simple Google search using the operator site:dunhamlaw.com "dunnam" would immediately reveal whether Defendants\u2019 website contains any reference to the "Dunnam" brand. Similarly, a Screaming Frog crawl\u2014a tool Brown claims to use\u2014would capture every piece of text content, metadata, and site architecture element on the domain. If any prohibited name formulation or brand-targeting existed, it would be plainly visible in such a crawl report. Brown provides no such report.')

doc.add_heading('C. The \u201cIf It\u2019s Visible to Google, It\u2019s Visible to You\u201d Principle', level=2)

add_para(doc, "A foundational principle of search engine optimization is that search engines can only index and rank content that is publicly accessible. If Defendants were using prohibited names on their website, embedding them in metadata, or targeting the \u201cDunnam & Dunnam\u201d brand through on-page content, this information would be visible to any person using a web browser, any search engine crawler, and any third-party SEO tool. The fact that Brown cannot produce hard evidence of any specific violation\u2014despite having access to the same public internet as everyone else\u2014strongly suggests that no such violation exists.")

# ============================================================
# V. METHODOLOGY
# ============================================================

doc.add_page_break()
doc.add_heading('V. Brown\u2019s Methodology Is Insufficient and Non-Reproducible', level=1)

doc.add_heading('A. Unnamed Tools and Undocumented Processes', level=2)

add_para(doc, 'Brown references "third-party SEO visibility data" and "third-party keyword visibility" throughout her report without ever identifying the specific tools used. Industry-standard practice for expert analysis requires naming tools (e.g., Ahrefs, SEMrush, SimilarWeb, Moz), specifying the exact date of data extraction, and providing the raw report outputs. All of these tools generate downloadable reports that serve as verifiable evidence. Brown provides none.')

doc.add_heading('B. Self-Identified Methodological Failures', level=2)

add_para(doc, "In Section XI of her report, Brown herself acknowledges that proper reproducibility would require documenting:")

add_bullet(doc, "Crawl scope and user-agent")
add_bullet(doc, "Date range")
add_bullet(doc, "Deduplication method for templated content")
add_bullet(doc, "Handling of dynamic and paginated content")
add_bullet(doc, "Phrase matching and variant handling logic")

add_para(doc, 'Brown admits she did not provide any of these parameters. This is a significant admission. Without a documented deduplication methodology, her keyword frequency counts (e.g., "450 mentions of \'Waco criminal attorney\'") are almost certainly inflated by templated content\u2014headers, footers, navigation menus, and sidebar elements that repeat across every page of a website. A single navigation element containing "Waco" that appears on 450 pages would be counted as "450 mentions" under an undeduplicated methodology, despite representing a single design choice.')

doc.add_heading('C. Qualitative Conclusions from Quantitative Claims', level=2)

add_para(doc, 'Brown admits that her January 2026 update is only a "qualitative comparison" to the August 2025 baseline, yet she draws the definitive conclusion that there has been "no material retreat" in restricted-area visibility. An expert cannot make quantitative claims ("no material retreat") based on admittedly qualitative analysis. Her approximate keyword counts, which she characterizes as "descriptive indicators" that "should not be treated as a definitive census," cannot support the categorical conclusions she draws elsewhere in the report.')

doc.add_heading('D. Publicly Verifiable Items Treated as Requiring Discovery', level=2)

add_para(doc, 'Brown\'s Verification/Proof Matrix (Attachment A) claims that several items are "not publicly verifiable" and require production from Defendants. In fact, many of these items are partially or fully verifiable through public tools:')

add_bullet(doc, " Fully verifiable through a Screaming Frog crawl, which Brown claims to use.", "Metadata cleanup:")
add_bullet(doc, " Verifiable through direct browser testing, HTTP header inspection, and redirect-checking tools.", "Domain redirects:")
add_bullet(doc, " Verifiable through the Google Ads Transparency Center, direct search observation in restricted counties, and third-party paid search intelligence tools (SEMrush, SpyFu, SimilarWeb).", "Paid ad presence:")
add_bullet(doc, " Verifiable through site crawling and content analysis.", "County targeting on-site:")

add_para(doc, "That Brown claims these items require discovery from Defendants, rather than conducting the public analysis herself, suggests either a lack of familiarity with the available tools or a recognition that such analysis would not support her conclusions.")

# ============================================================
# VI. ORGANIC RANKINGS
# ============================================================

doc.add_page_break()
doc.add_heading('VI. Organic Search Rankings Are Not Evidence of Misconduct', level=1)

doc.add_heading('A. How Search Engines Handle Similar Names', level=2)

add_para(doc, 'Brown\'s finding that DunhamLaw.com ranked for queries like "dunnam & dunnam" is presented as evidence of wrongdoing. It is not. Search engines, including Google, routinely surface results for approximate name matches, misspellings, and phonetically similar terms. When a user searches for "dunnam," Google may surface results for "dunham" because the algorithm recognizes the names as similar. This is core algorithmic behavior that neither party controls.')

add_para(doc, 'The closest Defendants could come to influencing rankings for the "Dunnam & Dunnam" brand would be to include that name in their website content, metadata, or link structure. This has never been the case. This can be verified by anyone through a simple site:dunhamlaw.com "dunnam" search, which would reveal any instance of the Dunnam brand appearing on Defendants\u2019 domain. No such evidence has been presented because none exists.')

doc.add_heading('B. Neither Party Controls Google\u2019s Algorithm', level=2)

add_para(doc, "It is a fundamental reality of modern search that no website operator controls how Google ranks their site for queries they have not specifically targeted. Google\u2019s ranking algorithm considers hundreds of factors, including domain name similarity, geographic proximity, practice area overlap, and user behavior signals. That two law firms with near-identical surnames, operating in overlapping geographic markets and the same practice areas, would appear in each other\u2019s branded search results is entirely expected and does not indicate any manipulation by either party.")

doc.add_heading('C. The \u201cAlgorithmic Ambiguity\u201d Argument Is Without Merit', level=2)

add_para(doc, 'Brown introduces the concept of "algorithmic ambiguity" and describes a chatbot incident where "incorrect or conflicting contact information" was surfaced. She admits this was "not consistently reproducible" and is "not relied upon as a standalone finding." This admission effectively concedes the finding has no evidentiary value. Defendants do not control third-party AI chatbot behavior, and isolated, non-reproducible algorithmic outputs are not evidence of misconduct.')

# ============================================================
# VII. PROACTIVE COMPLIANCE
# ============================================================

doc.add_heading('VII. Defendants Have Taken Proactive Steps Beyond the Judgment\u2019s Requirements', level=1)

add_para(doc, "Brown\u2019s report fails to give adequate weight to a significant fact: Defendants voluntarily redirect their primary domain DunhamLaw.com to DunhamJones.com for users located in the restricted counties.")

add_para(doc, "This is an extraordinary measure. Defendants operate a successful law practice with a 20+ year established domain. No provision of the 1998 Judgment requires them to redirect their primary domain for any geographic subset of users. That they have chosen to do so\u2014at the cost of diminished brand consistency and potential user confusion\u2014demonstrates a good-faith commitment to compliance that goes well beyond the Judgment\u2019s requirements.")

add_para(doc, 'In my professional experience managing enterprise-level web properties, geographic domain redirection is a technically sophisticated solution that requires ongoing maintenance and investment. It is the opposite of the "selective implementation choices" that Brown alleges.')

# ============================================================
# VIII. CORRECTIVE ADVERTISING
# ============================================================

doc.add_page_break()
doc.add_heading('VIII. Brown\u2019s Corrective Advertising Estimates Are Unsupported and Inflated', level=1)

doc.add_heading('A. No Methodology or Basis Provided', level=2)

add_para(doc, "Brown proposes annual corrective advertising costs of $65,000 to $120,000 per year for an unspecified multi-year period, without citing any methodology, market data, rate cards, comparable cases, or industry benchmarks. The figures appear to be arbitrary.")

doc.add_heading('B. Disproportionate to the Market', level=2)

add_para(doc, 'Waco, Texas (McLennan County) is a regional market with a metropolitan population of approximately 270,000. The proposed "paid search corrective ads" budget of $25,000\u2013$50,000 per year is orders of magnitude beyond what would be required for a legitimate search clarification campaign in a market of this size. For context, $50,000 in annual Google Ads spend in a regional legal market would purchase an extraordinarily dominant advertising presence\u2014far beyond any reasonable "corrective" purpose.')

doc.add_heading('C. Self-Serving Monitoring Fees', level=2)

add_para(doc, "Brown proposes $12,000\u2013$20,000 per year for \u201congoing monitoring and enforcement.\u201d At her stated rate of $125 per hour, this represents 96 to 160 hours of annual monitoring\u2014roughly 2 to 3 hours per week, every week, indefinitely. This is grossly disproportionate to the monitoring needs of a single regional trademark matter and appears designed to create a recurring revenue stream rather than address a genuine compliance need.")

doc.add_heading('D. Search Signals Do Not Require Multi-Year Correction', level=2)

add_para(doc, 'Brown asserts that search engine associations "do not self-correct simply because conduct stops" and that corrective efforts "must be sustained" over a "multi-year period." This is misleading. Search engine results and ranking signals are highly dynamic. Google recrawls and reindexes active websites on a continuous basis, typically within days to weeks. Historical ranking signals naturally decay as content is updated, removed, or superseded. A multi-year corrective campaign is not supported by how search engines actually function.')

doc.add_heading('E. Premature and Presumptive', level=2)

add_para(doc, "The entire corrective advertising analysis presumes that violations have been proven and damages have been established. Neither is the case. Calculating corrective advertising costs before establishing that any violation occurred puts the cart before the horse and should be disregarded.")

# ============================================================
# IX. QUALIFICATIONS
# ============================================================

doc.add_heading('IX. Observations on the Opposing Expert\u2019s Qualifications', level=1)

add_para(doc, "While the weight to be given expert testimony is a matter for the Court, I note the following regarding Brown\u2019s qualifications as they relate to the reliability of her opinions:")

add_bullet(doc, " not search marketing or advertising expertise.", "Brown\u2019s professional background prior to founding Digital Media Butterfly in 2012 was in system administration\u2014managing phone systems, hardware rollouts, and providing technical support for approximately 400 users. This is information technology work,")

add_bullet(doc, " in web design and development. She holds no degree in marketing, advertising, computer science, or any related field.", "Brown\u2019s formal education consists of coursework in a 1998 night program")

add_bullet(doc, " in the field of search marketing, SEO, or digital advertising.", "Brown has never previously testified as an expert witness and has no publications")

add_bullet(doc, "\u2014no enterprise, Fortune 500, or other significant engagements that would demonstrate experience with the kind of multi-domain, multi-market analysis she attempts in this report.", "Brown names no clients in her CV")

add_bullet(doc, " that exceed the scope of technical expertise and invade the province of the fact-finder. An expert in search marketing can opine on technical facts; she cannot opine on whether Defendants\u2019 \u201cselective implementation choices\u201d constitute intentional non-compliance.", "Brown\u2019s report makes legal conclusions about intent, willfulness, and compliance interpretation")

add_bullet(doc, "\u2014the same city as Plaintiff Dunnam & Dunnam\u2014which, while not dispositive, is relevant to evaluating independence in a local trademark dispute.", "Brown is located in Waco, Texas")

# ============================================================
# X. CONCLUSION
# ============================================================

doc.add_page_break()
doc.add_heading('X. Conclusion', level=1)

add_para(doc, "The Brown Report presents 15 pages of speculation, unsupported claims, and methodological shortcuts in an attempt to create the appearance of non-compliance where none has been demonstrated. Its core failings are:")

add_bullet(doc, " Brown does not produce a single screenshot of a prohibited advertisement, a single instance of a prohibited name on Defendants\u2019 website, or any other concrete artifact of non-compliance. If such evidence existed, it would be publicly visible and trivially easy to capture.", "No evidence of violation.")

add_bullet(doc, ' Brown treats the Judgment as a blanket prohibition on the use of the name "Dunham" in digital environments. It is not. The Judgment prohibits specific name formulations in specific counties\u2014a far narrower restriction than Brown applies.', "Misreading of the 1998 Judgment.")

add_bullet(doc, " Brown\u2019s analytical framework is structured so that no evidence could ever demonstrate compliance: if she finds something, it\u2019s a violation; if she finds nothing, it requires discovery to confirm. This is not rigorous analysis\u2014it is advocacy.", "Unfalsifiable reasoning.")

add_bullet(doc, " The corrective advertising estimates of $65,000\u2013$120,000 per year are unsupported by any methodology, wildly disproportionate to the regional market, and premature absent any established violation.", "Inflated damages.")

add_bullet(doc, " Brown gives no meaningful credit to Defendants\u2019 voluntary geographic domain redirection\u2014a significant, proactive compliance measure that goes beyond the Judgment\u2019s requirements.", "Ignored compliance efforts.")

add_para(doc, "In my professional opinion, Defendants\u2019 online properties do not demonstrate the violations alleged in the Brown Report. The report should be given little weight due to its methodological deficiencies, unsupported conclusions, and fundamental mischaracterization of both the applicable legal standard and the technical realities of search engine behavior.")

# ============================================================
# SIGNATURE
# ============================================================

add_para(doc, "")
add_para(doc, "The opinions expressed herein are based on my professional experience and analysis of publicly available information as described. I reserve the right to supplement this report upon receipt of additional materials or discovery.")

add_para(doc, "")
add_para(doc, "Executed on __________________, in Los Angeles, California.")

add_para(doc, "")
add_para(doc, "")
add_para(doc, "________________________________________")

p = doc.add_paragraph()
run = p.add_run("Kenny Hyder")
run.bold = True
p.paragraph_format.space_after = Pt(0)

p = doc.add_paragraph("Founder, Hyder Media")
p.paragraph_format.space_after = Pt(0)
p = doc.add_paragraph("kenny@hyder.me | (619) 850-7217")
p.paragraph_format.space_after = Pt(0)
p = doc.add_paragraph("hyder.me")

# Save
output = os.path.join(os.path.dirname(__file__), 'expert-rebuttal-report.docx')
doc.save(output)
print(f'Word document saved to {output}')
