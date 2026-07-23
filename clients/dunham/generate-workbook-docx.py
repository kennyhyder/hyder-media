#!/usr/bin/env python3
"""Generate Rebuttal Workbook as a Word document with Brown's claims and Kenny's notes."""

import json
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

h1 = doc.styles['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Pt(14)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
h1.paragraph_format.space_before = Pt(20)
h1.paragraph_format.space_after = Pt(8)

h2 = doc.styles['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Pt(11)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after = Pt(4)


def add_para(text, color=None, italic=False, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if color:
        run.font.color.rgb = color
    if italic:
        run.font.italic = True
    if bold:
        run.bold = True
    p.paragraph_format.space_after = Pt(6)
    return p


# Load notes
notes_path = os.path.join(os.path.dirname(__file__), 'rebuttal-notes-2026-02-27.json')
with open(notes_path, 'r') as f:
    notes = json.load(f)

# ============================================================
# TITLE PAGE
# ============================================================

logo_path = os.path.join(os.path.dirname(__file__), '..', 'assets', 'imgs', 'logos', 'hyder-media-logo.png')
if os.path.exists(logo_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(0.6))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Expert Rebuttal Workbook')
run.bold = True
run.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Caryn Brown, Digital Media Butterfly — Expert Report dated February 6, 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Dunnam & Dunnam LLP v. Dunham Law Firm, P.C., et al.')
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Civil Action No. 6:21-cv-1041-ADA-DGT')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('_' * 60)

add_para("")

# ============================================================
# 1998 JUDGMENT CONTEXT
# ============================================================

doc.add_heading('1998 Agreed Judgment — Key Terms', level=1)

add_para('Case: No. 98-1177-3, McLennan County, TX — 74th Judicial District, signed April 23, 1998.', bold=True)

add_para('Parties: Dunnam & Dunnam LLP (Plaintiffs) v. The Dunham Law Firm, Attorneys at Law, P.C., Paul J. Dunham, Jeffrey R. Casey, Tony Faulkner, and Bobby Hawkins (Defendants).')

add_para('Core restriction: Defendants permanently enjoined from using, directly or indirectly, in the following counties: Bell, McLennan, Coryell, Hill, Bosque, Johnson, Brazos, Hamilton, Somervell, Limestone, Freestone, Robertson, Falls, Madison, Navarro, and Leon — the names: "Dunham, Attorney at Law," "Dunhams, Attorney at Law," "Dunham Law Firm," "Dunham Firm," "Dunhams," "Dunham and Associates," and any variation using "Dunham" or "Dunnam" with the word "The" or "the," or with abbreviations (Inc., P.C., LLP), or names with the spelling "Dunnam" rather than "Dunham," or "Dunham" or "Dunnam" with a different given or surname.')

add_para('Advertising restriction: Defendants shall cancel all advertising or telephone listings using the specified names in the specified area, and shall not hereafter place any advertising or solicitation emphasizing "Dunham" or "Dunnam" in the specified area.')

p = add_para('Key limitation: The judgment restricts the use of specific names in specific counties. It does not restrict the defendants from practicing law or advertising their services in those counties under non-prohibited names (e.g., "Dunham & Jones").')
p.runs[0].bold = True

doc.add_page_break()

# ============================================================
# SECTIONS
# ============================================================

sections_data = [
    {
        "title": "I. Summary of Opinions",
        "ref": "Report pp. 1",
        "claims": 'Brown claims that defendants\' domains continue to show "ongoing exposure in restricted counties" and "continued algorithmic association with the Dunnam & Dunnam brand." She asserts the "net consumer-facing and algorithmic effect has not materially retreated" since August 2025, and that visibility has merely "shifted between DunhamLaw.com and DunhamJones.com" while the "aggregate footprint remains substantially similar."',
        "notes_key": "I. Summary of Opinions",
    },
    {
        "title": "II. Assignment and Scope",
        "ref": "Report pp. 1",
        "claims": 'Brown was retained by Dunnam & Dunnam\'s counsel to assess compliance with the 1998 Agreed Judgment and 2011 Rule 11 Agreement. Her scope covers domain use, redirections, SEO practices, organic rankings, paid search, geographic targeting, consumer/algorithmic confusion, and "technical feasibility of compliance." She states: "Because several key facts can be confirmed only from defendants\' internal systems... this report identifies those areas and refrains from assuming compliance absent documentary proof."',
        "notes_key": "II. Assignment and Scope",
    },
    {
        "title": "III. Sources, Methods, and Limitations",
        "ref": "Report pp. 2",
        "claims": 'Brown lists her sources as "publicly observable website content," "third-party SEO visibility data," and "Wayback Machine snapshots." Her methods include domain/site architecture review, third-party keyword visibility review, and examination of "patterns characteristic of technical SEO deployment." She acknowledges she lacks access to negative keywords, geo-exclusions, metadata exports, redirect inventories, and Search Console data. She states: "The absence of internal proof should not be treated as evidence of compliance."',
        "notes_key": "III. Sources, Methods, and Limitations",
    },
    {
        "title": "IV. Findings A — DunhamLaw.com",
        "ref": "Report pp. 3–4",
        "claims": 'A1 — Domain Use: DunhamLaw.com is active, "Dunham" appears in URLs/titles/headers/meta content, and the site "marketed services tied to restricted counties" with "Waco-related phrasing."\n\nA2 — Brand-Adjacent Visibility: DunhamLaw.com ranked for "dunnam & dunnam," "dunnam and dunnam," and "dunnam and dunnam Waco" in third-party organic data. She concedes paid advertising "cannot be verified" without account exports.\n\nA3 — Geographic Targeting: Site promoted services in McLennan County (Waco) and other restricted counties with "county mentions and Waco-specific phrasing at scale."\n\nA4 — Keyword Counts (Aug 2025): "Waco criminal attorney" — ~450 mentions; "Bell County" — ~195; "Brazos County" — ~70; "Coryell County" — ~65; "Hill County" — ~55; "Bosque County" — ~60; "Johnson County" — ~50; Other restricted counties — ~20–45 each.',
        "notes_key": "IV. Findings A \u2014 DunhamLaw.com",
    },
    {
        "title": "V. Findings B — DunhamJones.com",
        "ref": "Report pp. 5",
        "claims": 'B1 — Domain Timeline: DunhamJones.com registered June 17, 2011; first Wayback capture February 20, 2023. Brown suggests "restricted-county content may now be hosted on DunhamJones.com rather than DunhamLaw.com."\n\nB2 — Content Migration: Third-party data shows DunhamJones.com ranks for "Waco-specific and other restricted-county criminal defense queries." She calls this "content migration or re-optimization rather than removal."\n\nB3 — Paid Search: No paid keyword activity detected for DunhamJones.com, but Brown argues this "does not establish compliance" because campaigns "can be paused/reactivated."',
        "notes_key": "V. Findings B \u2014 DunhamJones.com",
    },
    {
        "title": "VI. Developments Since August 2025",
        "ref": "Report pp. 6",
        "claims": 'Brown asserts "no material retreat from restricted-area visibility or brand-adjacent search exposure at the aggregate level" between August 2025 and January 2026. She claims that any visibility shift between the two domains means the "net consumer-facing and algorithmic effect remains substantively similar." She states: "Claims of meaningful corrective action since August 2025 cannot be verified from public observation alone."',
        "notes_key": "VI. Developments Since August 2025",
    },
    {
        "title": "VII. Intent, Ability to Comply, and Post-Notice Conduct",
        "ref": "Report pp. 7",
        "claims": 'A — Intent: Because defendants manage their sites, they have "technical capacity" to comply, and any non-compliance "is not attributable to technical impossibility" but rather "selective implementation choices."\n\nB — Post-Notice: Defendants continued modifying their digital presence after being on notice, but "core elements... have not been conclusively cured." She calls this "relevant evidence because it shows compliance steps were feasible but not completed."\n\nC — Algorithmic/AI Confusion: A chatbot on defendants\' website "surfaced incorrect or conflicting contact information associated with Dunnam & Dunnam." Brown admits this "was not consistently reproducible" and is "not relied upon as a standalone finding."',
        "notes_key": "VII. Intent, Ability to Comply, and Post-Notice Conduct",
    },
    {
        "title": "VIII. Technical Simplicity of Compliance",
        "ref": "Report pp. 8",
        "claims": 'Brown asserts that compliance is "technically straightforward and comparatively inexpensive," listing: disabling DNS for "prohibited domains," deactivating/redirecting those domains, removing prohibited references from content/metadata/structured data, Search Console URL inspection and re-index requests, and negative keyword implementation with geo-exclusions in paid channels.',
        "notes_key": "VIII. Technical Simplicity of Compliance",
    },
    {
        "title": "IX–X. Verification Matrix & Recommended Discovery",
        "ref": "Report pp. 9–11",
        "claims": 'Brown presents a "Verification/Proof Matrix" showing what she claims is and isn\'t publicly verifiable, and requests production of: full site backups, keyword/metadata audits, Google Ads exports, Analytics/Search Console data, SEO vendor communications, and full domain inventories. She states: "The absence of verification is not evidence of compliance. It reflects the limits of public observation and reinforces the need for documentary proof."',
        "notes_key": "IX\u2013X. Verification Matrix & Recommended Discovery",
    },
    {
        "title": "XI. Quantitative Indicators & Methodological Transparency",
        "ref": "Report pp. 12",
        "claims": 'Brown retroactively acknowledges that her August 2025 keyword counts were "approximate values" from "crawls and text analysis" and that the January 2026 update is only a "qualitative comparison." She admits that proper reproducibility would require documenting crawl scope, user-agent, date range, deduplication, dynamic content handling, and phrase matching logic — none of which she provided.',
        "notes_key": "XI. Quantitative Indicators & Methodological Transparency",
    },
    {
        "title": "XIII. Corrective Advertising",
        "ref": "Report pp. 13–14",
        "claims": 'Brown proposes multi-year corrective advertising to "unwind historical associations," arguing that search engine associations "do not self-correct simply because conduct stops." She proposes annual costs of: Paid search corrective ads ($25,000–$50,000/year), SEO corrective content & signal repair ($20,000–$35,000/year), Directory/platform remediation ($8,000–$15,000/year), Ongoing monitoring & enforcement ($12,000–$20,000/year). Total: $65,000–$120,000/year for multiple years.',
        "notes_key": "XIII. Corrective Advertising",
    },
    {
        "title": "XIV. Expert Qualifications & Credibility",
        "ref": "Report pp. 15 + CV",
        "claims": 'Brown\'s qualifications per her CV and filing: Owner, Digital Media Butterfly (est. 2012, Waco, TX). Prior role: System Administrator at Brazos Higher Education (managing phone systems, 400 users, hardware rollouts). Education: "Formal coursework in web design and development (night program, 1998)." No prior expert testimony (stated in filing). No publications. No named clients. No industry conference speaking at recognized events. Rate: $125/hour. Tools: Ahrefs, Screaming Frog, Wayback Machine, WHOIS.',
        "notes_key": "XIV. Expert Qualifications & Credibility",
    },
]

for sec in sections_data:
    doc.add_heading(sec["title"], level=1)

    # Reference
    p = doc.add_paragraph()
    run = p.add_run(sec["ref"])
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    run.font.italic = True

    # Brown's claims
    doc.add_heading("Brown's Claims", level=2)
    for para_text in sec["claims"].split('\n\n'):
        p = doc.add_paragraph()
        run = p.add_run(para_text.strip())
        run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)
        p.paragraph_format.space_after = Pt(6)

    # My notes
    doc.add_heading("Rebuttal Notes", level=2)
    note_text = notes.get(sec["notes_key"], "")
    if note_text:
        for para_text in note_text.split('\n\n'):
            cleaned = para_text.strip()
            if cleaned:
                p = doc.add_paragraph()
                run = p.add_run(cleaned)
                p.paragraph_format.space_after = Pt(6)
    else:
        add_para("[No notes entered]", color=RGBColor(0x9c, 0xa3, 0xaf), italic=True)

# Save
output = os.path.join(os.path.dirname(__file__), 'rebuttal-workbook.docx')
doc.save(output)
print(f'Word document saved to {output}')
